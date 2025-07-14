import os
import time
import tkinter as tk
import customtkinter as ctk
import PyPDF2
import pandas as pd
import sys
import re
from comtypes.client import CreateObject
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageOps
from pathlib import Path
from tkinter import messagebox
from tkcalendar import Calendar
from datetime import datetime
from tkinter import filedialog
from logic.variables import *
from docxtpl import DocxTemplate
from logic.utilities import *
from ui.styles import *
from .variables import *

variables_entries = {}
variables_combo= {}

# Lista para almacenar los datos ingresados
datos = []
# Variable para almacenar la ruta seleccionada

def crear_popup(title, message, close=True):
    popup= tk.Toplevel()
    icon_path = resource_path('assets/icons/logo_santander_ico.ico')
    popup.iconbitmap(icon_path)
    popup.title(title)
    popup.geometry("320x190")
    popup.resizable(False, False)
    popup.configure(bg=font_gray_color)
    text_widget = tk.Text(popup, bg=font_gray_color, fg=font_light_color, bd=0, height=4)
    text_widget.pack(padx=10, pady=10)
    text_widget.insert(tk.END, message)
    close_button = ctk.CTkButton(popup, text="Cerrar", command=popup.withdraw, fg_color=options_button_color, width=button_popup_width, height=button_popup_height, font=mi_fuente_button, text_color= font_dark_color, corner_radius= corner_button_radius, cursor=cursor_button_type, hover_color= options_button_hover_color)
    close_button.pack(pady=10)
    if not close:
        close_button.pack_forget()
    return popup

def crear_combobox(frame,texto,labeltexto,opciones,opcselec):
    frame.rowconfigure(0,weight=1)
    frame.columnconfigure(0, weight=1)
    frame.columnconfigure(1, weight=1)
    label = tk.Label(frame, bg=font_gray_color,text=labeltexto,width=20,anchor="w", font=mi_fuente_little, fg=font_light_color)
    label.grid(row=0, column=0, padx=5, pady=10, sticky="nswe")

    combo = ctk.CTkComboBox(frame, values=opciones,variable=opcselec,state="readonly", width=135, text_color=font_dark_color,corner_radius= corner_button_radius, fg_color=font_light_color,border_color=font_gray_color, dropdown_fg_color=font_light_color, dropdown_text_color=font_dark_color, font = ("Courier", 14, "bold"), dropdown_font=("Courier", 14, "bold"), dropdown_hover_color=options_button_color)
    combo.grid(row=0, column=1, padx=5, pady=10, sticky="nswe")
    variables_combo[texto] = opcselec

# Función para crear un par de labels y entrys con alineación uniforme
def crear_campos(frame,texto,labeltexto):
    frame.rowconfigure(0,weight=1)
    frame.columnconfigure(0, weight=1)
    frame.columnconfigure(1, weight=1)
    label = tk.Label(frame, text=labeltexto, width=20, anchor="w",bg=font_gray_color, fg=font_light_color, font=mi_fuente_little)
    label.grid(row=0, column=0, padx=5, pady=10, sticky="nswe")

    # Variable asociada al Entry
    var_entry = tk.StringVar()
    entry = ctk.CTkEntry(frame, textvariable=var_entry,width=135, height=25, fg_color=font_light_color, bg_color=font_gray_color, font=mi_fuente_little, corner_radius=corner_button_radius, border_color=font_light_color)
    entry.grid(row=0, column=1, padx=5, pady=10, sticky="nswe")

    # Guardar la variable asociada al texto del campo en el diccionario
    variables_entries[texto] = var_entry

# Función para actualizar la tabla con los datos ingresados
def actualizar_tabla(tabla):
    # Obtener las filas actuales en la tabla
    filas_actuales = tabla.get_children()

    # Eliminar solo las filas que ya no están en la lista 'datos'
    for i, fila in enumerate(datos):
        if i >= len(filas_actuales):
            tabla.insert("", "end", values=fila)
        else:
            tabla.item(filas_actuales[i], values=fila)

    # Ajustar la tabla si hay menos filas en 'datos' que en la tabla
    if len(filas_actuales) > len(datos):
        for i in range(len(datos), len(filas_actuales)):
            tabla.delete(filas_actuales[i])


def generar_documentos(tabla):
    title = None
    message = None
    popup_generando = None
    try:
        # Verificar si la tabla tiene al menos una fila con todos los campos completos
        if not tabla.get_children():
            # Si la tabla está vacía
            title = "Error"
            message = "La tabla está vacía."
            crear_popup(title, message, close=True)
            return
        
        print("Generando documentos de evidencias...")
        popup_generando = crear_popup("Generador de evidencias", "Generando documentos de evidencias...", close=False)
        popup_generando.update_idletasks()

        # Limpia la carpeta de evidencias antes de generar los documentos
        crear_carpeta_evidencias()
        
        base_dir = Path(__file__).parent if "__file__" in locals() else Path.cwd()
        evidencias_dir = base_dir / "EVIDENCIAS"  # Directorio de evidencias
        # Crear la carpeta de evidencias si no existe
        evidencias_dir.mkdir(parents=True, exist_ok=True)
        print(f"Directorio de evidencias: {evidencias_dir}")

        documento_word = Path(resource_path('template/Plantilla.docx'))
        for item in tabla.get_children():
            fila = tabla.item(item, "values")
            print(f"Procesando fila: {fila}")

            doc = DocxTemplate(documento_word)  # Cargar la plantilla de Word
            context = {}  # Crear un diccionario para los datos de la plantilla
            # Iterar sobre los datos de la fila y agregarlos al contexto
            for nombre_campo, dato in zip(variables_entries.keys(), fila):
                print(nombre_campo, dato)
                context[nombre_campo] = dato
            for i, columna in enumerate(campos_tabla):
                context[columna] = fila[i]
            
            print(context)
            doc.render(context)
            print(f"Contexto renderizado: {context}")

            if len(fila) >= 2:
                ruta_imagenes = fila[-1]  # Obtener la ruta de las imágenes desde la tabla
                print(f"Ruta de imágenes: {ruta_imagenes}")
                # Agregar las imágenes al documento
                add_images_to_doc(doc, ruta_imagenes)
                time.sleep(1)

            nombredocumento = "Reporte de Evidencias - " + datetime.now().strftime("%d-%m-%Y %H-%M-%S")
            # Guardar el documento generado en formato docx
            docx_path = evidencias_dir / f"{nombredocumento}.docx"
            doc.save(docx_path)
            print(f"Documento guardado en: {docx_path}")
            # Convertir el documento a PDF
            pdf_path = evidencias_dir / f"{nombredocumento}.pdf"
            word = CreateObject('Word.Application')
            doc = word.Documents.Open(str(docx_path))
            doc.SaveAs(str(pdf_path), FileFormat=17)  # 17 is the code for wdFormatPDF
            doc.Close()
            word.Quit()
            print(f"Documento convertido a PDF en: {pdf_path}")
            time.sleep(1)

        # Combinar todos los PDF generados en un solo archivo
        # merge_pdfs_in_folder(evidencias_dir, "Reporte de Evidencias - Todos.pdf")

        popup_generando.destroy()
        
        title = "Generar evidencias"
        message = "Los documentos se han generado correctamente."

    except Exception as e:
        title = "Error"
        message = f"Ocurrió un error durante la generación de documentos: {str(e)}"
        crear_popup(title, message, close=True)
        print(f"Error: {str(e)}")
        return

    finally:
        if popup_generando:
            popup_generando.destroy()

    popup = crear_popup(title, message, close=False)
    close_button = ctk.CTkButton(popup, text="Ver archivos", command=lambda: abrir_carpeta_evidencias(popup), fg_color=options_button_color, width=button_popup_width, height=button_popup_height, font=mi_fuente_button, text_color=font_dark_color, corner_radius=corner_button_radius, cursor=cursor_button_type, hover_color=options_button_hover_color)
    close_button.pack(pady=10)


def adjust_image_size(image_path):
    with Image.open(image_path) as img:
        width, height = img.size
        # Convert pixel dimensions to cm (assuming 96 DPI)
        width_cm = width / 96 * 2.54
        height_cm = height / 96 * 2.54

        if height_cm > width_cm:
            new_height_cm = 10
            aspect_ratio = width_cm / height_cm
            new_width_cm = new_height_cm * aspect_ratio
        else:
            new_width_cm = 15
            aspect_ratio = height_cm / width_cm
            new_height_cm = new_width_cm * aspect_ratio

        return new_width_cm, new_height_cm
    
def add_border_to_image(image_path, border_size_cm):
    with Image.open(image_path) as img:
        # Calculate the border size in pixels (assuming 96 DPI)
        border_size_px = int(border_size_cm * 96 / 2.54)
        img_with_border = ImageOps.expand(img, border=border_size_px, fill='black')
        return img_with_border
    
def add_images_to_doc(doc, folder_path):
    sorted_files = []
    enumerate_files = []

    for filename in os.listdir(folder_path):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.pcx', '.tga')):
            if re.search(r'^\d+', filename):
                enumerate_files.append(filename)
            else:
                sorted_files.append(filename)

    # Sort enumerated files by leading number
    enumerate_files.sort(key=lambda x: int(re.match(r'^\d+', x).group()))

    # Sort files by modification date (oldest to newest)
    sorted_files.sort(key=lambda x: os.path.getmtime(os.path.join(folder_path, x)))

    # Combine both lists
    enumerate_files.extend(sorted_files)

    for filename in enumerate_files:
        image_path = os.path.join(folder_path, filename)
        new_image = add_border_to_image(image_path, 0.1)
        temp_image_path = os.path.join(folder_path, 'temp_' + filename)
        new_image.save(temp_image_path)

        width_cm, height_cm = adjust_image_size(temp_image_path)

        # Remove the file extension from the filename
        image_name = os.path.splitext(filename)[0]

        # Add image name as a paragraph with left indentation
        paragraph = doc.add_paragraph(image_name)
        paragraph.paragraph_format.left_indent = Cm(1)  # Adjust the indentation as needed
        paragraph.paragraph_format.keep_with_next = True  # Keep the title with the image

        # Add an empty paragraph to create a line break and keep it with the next paragraph
        empty_paragraph = doc.add_paragraph()
        empty_paragraph.paragraph_format.keep_with_next = True

        # Add image to the document and center it
        paragraph_img = doc.add_paragraph()
        paragraph_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph_img.add_run()
        run.add_picture(temp_image_path, width=Cm(width_cm), height=Cm(height_cm))

        # Add an empty paragraph to create space between images
        doc.add_paragraph()

        # Remove temporary image
        os.remove(temp_image_path)

def merge_pdfs_in_folder(folder_path, output_filename):
    # Create the "todos" folder inside the folder_path
    todos_folder_path = os.path.join(folder_path, 'todos')
    os.makedirs(todos_folder_path, exist_ok=True)
    
    # Update the output file path to be inside the "todos" folder
    output_file_path = os.path.join(todos_folder_path, output_filename)
    
    merger = PyPDF2.PdfMerger()

    pdf_files = [os.path.join(folder_path, file) for file in os.listdir(folder_path) if file.endswith('.pdf')]

    for pdf_file in pdf_files:
        merger.append(pdf_file)

    with open(output_file_path, 'wb') as output_file:
        merger.write(output_file)

    merger.close()

def abrir_carpeta_evidencias(popup):
    popup.destroy()
    base_dir = Path(__file__).parent if "__file__" in locals() else Path.cwd()
    evidencias_dir = base_dir / "EVIDENCIAS"
    os.startfile(evidencias_dir)

def crear_carpeta_evidencias():
    base_dir = Path(__file__).parent if "__file__" in locals() else Path.cwd()
    evidencias_dir = base_dir / "EVIDENCIAS"
    # Verificar si la carpeta de evidencias existe
    if evidencias_dir.exists():
        # Crear la carpeta de evidencias nuevamente
        evidencias_dir.mkdir(parents=True, exist_ok=True)

def generar_excel(tabla):
    try:
        # Obtener los datos de la tabla
        datos_tabla = []
        for fila_id in tabla.get_children():
            datos_fila = tabla.item(fila_id)["values"]
            # Completar las columnas faltantes con valores nulos
            while len(datos_fila) < 7:
                datos_fila.append(None)
            datos_tabla.append(datos_fila)

        if not datos_tabla:
            return

        # Crear un DataFrame a partir de los datos de la tabla
        df = pd.DataFrame(datos_tabla, columns=["CICLO", "ANALISTA", "CASOPRUEBA",  "PROYECTO", "FECHA", "ESTADO", "IMG"])

        # Obtener la ruta para guardar el archivo Excel
        ruta_guardar = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])

        if ruta_guardar:
            # Guardar el DataFrame como un archivo Excel
            df.to_excel(ruta_guardar, index=False)

            popup = crear_popup("Éxito", "El archivo Excel se ha generado correctamente.", close=False)
            open_button = ctk.CTkButton(popup, text="Ver archivo", command = lambda: abrir_archivo_excel(ruta_guardar, popup), fg_color=options_button_color, width=button_popup_width, height=button_popup_height, font=mi_fuente_button, text_color= font_dark_color, corner_radius= corner_button_radius, cursor=cursor_button_type, hover_color= options_button_hover_color)
            open_button.pack(pady=10)
          
    except Exception as e:
        messagebox.showerror("Error", f"Ocurrió un error durante la generación del archivo Excel: {str(e)}")
        print(f"Ocurrió un error durante la generación del archivo Excel: {str(e)}")

def abrir_archivo_excel(ruta_guardar, popup):
    os.startfile(ruta_guardar)
    popup.destroy()

def limpiar_tabla(tabla):
    # Limpiar la tabla
    for i in tabla.get_children():
        tabla.delete(i)

# Inicializar un diccionario para almacenar el ancho máximo de cada columna
max_widths = {campo: 0 for campo in campos_tabla}

def calcular_ancho(texto):
    return len(texto) * 6  # Aproximadamente 6 píxeles por carácter

def insertar_datos(tabla, datos):
    for i, fila in enumerate(datos):
        tabla.insert('', 'end', values=fila)
        for j, dato in enumerate(fila):
            ancho_dato = calcular_ancho(str(dato))
            # Actualizar el ancho máximo si el ancho del dato es mayor
            if ancho_dato > max_widths[campos_tabla[j]]:
                max_widths[campos_tabla[j]] = ancho_dato
                tabla.column(campos_tabla[j], minwidth= 100, width=ancho_dato)

def cargar_desde_excel(tabla):
    try:
        # Abrir el archivo de Excel
        ruta_excel = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if ruta_excel:
            # Leer el archivo de Excel
            df = pd.read_excel(ruta_excel, keep_default_na=False)
            limpiar_tabla(tabla)
            insertar_datos(tabla, df.values.tolist())
    except Exception as e:
        messagebox.showerror("Error", f"Ocurrió un error durante la carga desde Excel: {str(e)}")

def borrar_fila(tabla, boton_actualizar, boton_borrar_fila):
    
    boton_actualizar.grid_forget() # Obtener la fila seleccionada
    boton_borrar_fila.grid_forget() # Ocultar el botón de "Actualizar"
    seleccion = tabla.selection()
    # Verificar si se ha seleccionado una fila
    if seleccion:
        # Eliminar la fila seleccionada de la tabla
        tabla.delete(seleccion)
        # Eliminar los datos asociados de la lista 'datos'
        for item in seleccion:
            # Obtener el índice de la fila en 'datos'
            index = int(tabla.index(item))
            # Eliminar la fila de 'datos'
            del datos[index]

def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

# Función para seleccionar una ruta
def seleccionar_ruta(ruta_seleccionada):
    ruta = filedialog.askdirectory()
    ruta_seleccionada.set(ruta)

def seleccionar_fecha(root, var_entry):
    top = tk.Toplevel(root)
    top.configure(bg=font_dark_color)
    top.title("Seleccionar Fecha")
    top.geometry("350x300")
    top.resizable(False, False)
    icon_path = resource_path('assets/icons/logo_santander_ico.ico')
    top.iconbitmap(icon_path)
    cal = Calendar(top, selectmode="day", date_pattern="dd/mm/yyyy", maxdate=datetime.today(),bg=font_dark_color, fg=font_light_color, font=mi_fuente_little, locale="es_CO", selectbackground=iniciar_button_color)
    cal.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")

    def seleccionar():
        fecha = cal.get_date()
        var_entry.set(fecha)
        top.destroy()

    boton_seleccionar = ctk.CTkButton(top, text="Seleccionar", command=seleccionar, fg_color=iniciar_button_color, width=ctk_button_width, height=ctk_button_height, font=mi_fuente_button, text_color= font_dark_color, corner_radius= corner_button_radius, cursor=cursor_button_type, hover_color= options_button_hover_color)
    boton_seleccionar.grid(row=1, column=0, padx=10, pady=10)